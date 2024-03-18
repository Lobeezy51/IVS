import sqlite3
import datetime
import os
import streamlit as st
from docx import Document
import base64
import uuid

# Define the path to your SQLite database file
db_path = rf'C:\sqlite3\BoresightLog_240226.db'

# Define the function to execute the query and print the results
def execute_and_print_all(parent_serial_numbers, db_path):
    all_serial_dict = {}  # Initialize the dictionary
    part_numbers = {
        '90095047', '90095191-01', '90095045-01', '90095189-02',
        '90095191-04', '90095191-03', '90095191-02', '90095191-05',
        '90095191-06', '90095045-02', '90095189-03', '90095194',
        '90095189-01', '90095045-07', '90095191-08', '90095191-09',
        '90095191-10', '90095191-11', '90095191-12', '90095045-08', '90095193',
    }

    serial_patterns = {
        '[HO90095193][1.0]': '[HO90095193][0.0]',
        '[ZO90095193][1.0]': '[ZO90095193][0.0]',
        '[HO90095191-02][4.0]': '[HO90095191-04][4.0]',
        '[ZO90095191-02][4.0]': '[ZO90095191-04][4.0]',
        '[HO90095191-04][2.0]': '[HO90095191-02][2.0]',
        '[ZO90095191-04][2.0]': '[ZO90095191-02][2.0]',
        '[HO90095047][1.0]': '[HO90095047][0.0]',
        '[ZO90095047][1.0]': '[ZO90095047][0.0]',
        '[HO90095194][0.0]': '[HO90095194][7.0]',
        '[ZO90095194][0.0]': '[ZO90095194][7.0]',
    }

    def execute_and_print(part_number, parent_serial_number, cursor):
        example_query = """SELECT ParentSerialNumber, SerialNumber, Channel FROM BoresightLog_240226
                       WHERE ParentSerialNumber=? AND PartNumber=?;"""
        cursor.execute(example_query, (parent_serial_number, part_number))

        results = cursor.fetchall()

        serial_dict = {}

        for row in results:
            if parent_serial_number == optical_box_ho_entry:
                original_serial_number = f'[HO{part_number}][{row[2]}]'
            else:
                original_serial_number = f'[ZO{part_number}][{row[2]}]'

            for pattern, replacement in serial_patterns.items():
                if pattern in original_serial_number:
                    updated_serial_number_0 = original_serial_number
                    updated_serial_number_1 = original_serial_number.replace(pattern, replacement)
                else:
                    updated_serial_number_0 = original_serial_number
                    updated_serial_number_1 = original_serial_number

                serial_dict[updated_serial_number_0] = row[1]
                serial_dict[updated_serial_number_1] = row[1]

        #for key, value in serial_dict.items():
            #st.write(f"'{key}': {value}")

        return serial_dict  # Return the serial_dict

    with sqlite3.connect(db_path) as sqlite_connection:
        cursor = sqlite_connection.cursor()

        for parent_serial_number in parent_serial_numbers:
            for part_number in part_numbers:
                serial_dict = execute_and_print(part_number, parent_serial_number, cursor)
                all_serial_dict.update(serial_dict)  # Update the accumulated dictionary

    sqlite_connection.commit()

    return all_serial_dict  # Return the accumulated serial_dict

def update_document():
    # Fetch serial data using the new function
    optical_box_ho_value = optical_box_ho_entry
    optical_box_zo_value = optical_box_zo_entry
    parent_serial_numbers = [optical_box_ho_value, optical_box_zo_value]
    serial_dict = execute_and_print_all(parent_serial_numbers, db_path)

    # Validate P-Number
    if not p_number_entry.isdigit():
        st.error("Invalid P-Number. Please enter a numeric value.")
        return

    # Validate Fiber Bundle
    if not fiber_bundle_entry.isdigit():
        st.error("Invalid Fiber Bundle. Please enter a numeric value.")
        return

    # Validate HO SN
    if not optical_box_ho_value.isdigit():
        st.error("Invalid HO SN. Please enter a numeric value.")
        return

    # Validate ZO SN
    if not optical_box_zo_value.isdigit():
        st.error("Invalid ZO SN. Please enter a numeric value.")
        return
    
    # Validate Technician's Pin
    if not technician_pin_entry.isdigit():
        st.error("Invalid Technician's Pin. Please enter a numeric value.")
        return

    # Specify the paths for the original and updated documents
    document_path = os.path.join(rf"C:\\Users\\tony.loberg\\Pictures\\Python", "verifyModule_P-.docx")
    updated_document_name = os.path.join(rf"C:\\Users\\tony.loberg\\Pictures\\Python", f"verifyModule_P{p_number_entry}.docx")

    # Load the original document
    doc = Document(document_path)

    replacement_dict = {
        '[USERINTIALS]': technician_pin_entry,
        '[PNUMBER]': p_number_entry,
        '[FIBERBUNDLE]': fiber_bundle_entry,
        '[HOSN]': optical_box_ho_value,
        '[ZOSN]': optical_box_zo_value,
        '[INPUTFIBERHO]': 'input_fiber_ho_entry',
        '[INPUTFIBERZO]': 'input_fiber_zo_entry',
        '[TODAY]': datetime.datetime.today().strftime("%Y%m%d_%H%M"),
    }

    # Add serial data to replacement_dict
    replacement_dict.update(serial_dict)

    # Iterate through tables, rows, cells, paragraphs, and runs in the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for placeholder, value in replacement_dict.items():
                            run.text = run.text.replace(placeholder, value)

    # Save the updated document
    try:
        os.remove(updated_document_name)
    except FileNotFoundError:
        pass

    doc.save(updated_document_name)

    # Check if the document exists and is non-empty
    if os.path.exists(updated_document_name) and os.path.getsize(updated_document_name) > 0:
        return updated_document_name
    else:
        st.error("Error: Updated document not found or empty.")
        return None

# Interface
st.set_page_config(layout="wide", initial_sidebar_state="collapsed")

st.title('Integration Data Entry Form')

# User Information Frame
st.subheader("User Information")
user_container = st.container(border=True)
technician_pin_entry = user_container.text_input("Technician's Pin")

# Unit Information Frame
st.subheader("Unit Information")
unit_container = st.container(border=True)
col1, col2 = unit_container.columns(2)
p_number_entry = col1.text_input("P-Number")
fiber_bundle_entry = col1.text_input("Fiber Bundle")
optical_box_ho_entry = col2.text_input("HO SN")
optical_box_zo_entry = col2.text_input("ZO SN")

# Update Document Button
if st.button("Update Document"):
    updated_document_path = update_document()
    if updated_document_path:
        st.success("Document Updated Successfully!")
        
        # Display the updated Word document
        with open(updated_document_path, "rb") as file:
            doc_bytes = file.read()
            b64 = base64.b64encode(doc_bytes).decode()
            button_label = "Download Updated Document"
            button_uuid = str(uuid.uuid4())  # Generate a random UUID
            custom_css = f"""<style>
                            #{button_uuid} {{
                                width: auto;
                                display: inline-block;
                                text-align: center;
                            }}
                            </style>"""
            st.markdown(custom_css, unsafe_allow_html=True)
            href = f'<a href="data:application/octet-stream;base64,{b64}" download="verifyModule_P{p_number_entry}.docx"><button id="{button_uuid}">{button_label}</button></a>'
            st.markdown(href, unsafe_allow_html=True)
